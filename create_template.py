"""
Создаёт шаблон Word документа заявки.
Запусти один раз: python create_template.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:{}'.format(edge)
        element = OxmlElement(tag)
        element.set(qn('w:val'), kwargs.get(edge, 'single'))
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        tcBorders.append(element)
    tcPr.append(tcBorders)

def bold_run(para, text, size=10):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return run

def normal_run(para, text, size=10):
    run = para.add_run(text)
    run.font.size = Pt(size)
    return run

doc = Document()

# Поля страницы
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# Заголовок (правый угол)
header_para = doc.add_paragraph()
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = header_para.add_run("Приложение №01 к договору № ARGT-03-14/2026\nоб организации автомобильных перевозок\nот 03 марта 2026 года")
r.font.size = Pt(10)
r.bold = True

doc.add_paragraph()

# Номер заявки
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
bold_run(title, "Заявка №{{ZAYAVKA_NUM}}", size=12)

# Город и дата
city_para = doc.add_paragraph()
city_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
normal_run(city_para, "г. Астана                                                                                              {{DATE}}")

doc.add_paragraph()

# Вводный текст
intro = doc.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal_run(intro, "В соответствии с Договором об организации автомобильных перевозок №ARGT-03-14/2026 от 03 марта 2026г. Заказчик просит Перевозчика оказать услуги по перевозке грузов Заказчика на следующих условиях:")

doc.add_paragraph()

# Секция 1: Заполняется Заказчиком
sec1 = doc.add_paragraph()
r = sec1.add_run("Заполняется Заказчиком:")
r.bold = True
r.underline = True
r.font.size = Pt(10)

# Таблица 1
t1_data = [
    ("Заказчик", "ТОО «Astana Railway Group Trading»"),
    ("Грузоотправитель", "ТОО «Astana Railway Group Trading»"),
    ("Грузополучатель", "{{GRUZOPOLUCHATEL}}"),
    ("Маршрут перевозки", "{{MARSHRUT}}"),
    ("Тип автомашины и количество автомашин", "{{TIP_AVTO}}"),
    ("Тип груза (наименование, размер, вес, упаковка),\nособенности и характеристика груза", "{{TIP_GRUZA}}"),
    ("Температура (необходимая температура для перевозки)", "{{TEMPERATURA}}"),
    ("Заявленная стоимость груза", "{{STOIMOST_GRUZA}}"),
    ("Погрузка (адрес, дата и время)", "{{POGRUZKA}}"),
    ("Контактное лицо на погрузке (Ф.И.О, телефон)", "{{KONTAKT_POGRUZKA}}"),
    ("Разгрузка (адрес, дата и время)", "{{RAZGRUZKA}}"),
    ("Контактное лицо на разгрузке (Ф.И.О, телефон)", "{{KONTAKT_RAZGRUZKA}}"),
    ("Условия погрузки/выгрузки (боковая, задняя, верхняя)", "{{USLOVIYA_POGRUZKI}}"),
    ("Страхование груза", "{{STRAKHOVANIE}}"),
    ("Особые условия", "{{OSOBYE_USLOVIYA}}"),
]

table1 = doc.add_table(rows=len(t1_data), cols=2)
table1.style = 'Table Grid'
for i, (label, value) in enumerate(t1_data):
    row = table1.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    # Ширина колонок
    row.cells[0].width = Cm(8)
    row.cells[1].width = Cm(9.5)

doc.add_paragraph()

# Секция 2: Международные перевозки
sec2 = doc.add_paragraph()
r = sec2.add_run("Заполняется Заказчиком при запросе на перевозки в международном сообщении:")
r.bold = True
r.underline = True
r.font.size = Pt(10)

t2_data = [
    ("Код ТНВЭД", "{{KOD_TNVED}}"),
    ("Сторона ответственная за таможенное оформление в пункте погрузки", "ТОО «Astana Railway Group Trading»"),
    ("Сторона ответственная за таможенное оформление в пункте разгрузки", "{{TAMOZH_RAZGRUZKA}}"),
    ("Условия простоя в местах погрузки-выгрузки (таможенного оформления)", "В соответствии с условиями договора"),
    ("Ответственность Заказчика за простой транспортного средства в месте погрузки и разгрузке сверх установленных Договором сроков за каждые сутки простоя", "В соответствии с условиями договора"),
]

table2 = doc.add_table(rows=len(t2_data), cols=2)
table2.style = 'Table Grid'
for i, (label, value) in enumerate(t2_data):
    row = table2.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[0].width = Cm(8)
    row.cells[1].width = Cm(9.5)

doc.add_paragraph()

# Секция 3: Заполняется Перевозчиком
sec3 = doc.add_paragraph()
r = sec3.add_run("Заполняется Перевозчиком:")
r.bold = True
r.underline = True
r.font.size = Pt(10)

t3_data = [
    ("Стоимость услуг Перевозчика", "{{STOIMOST_PEREVOZCHIKA}}"),
    ("Марка машины, гос. номер и номер п/п", "{{MARKA_MASHINY}}"),
    ("Ф.И.О. водителя, паспортные данные, контактные данные", "ИИН {{IIN}}\nВыдано МВД РК от {{VYDANO}}\nТел.: {{TEL_VODITEL}}"),
]

table3 = doc.add_table(rows=len(t3_data), cols=2)
table3.style = 'Table Grid'
for i, (label, value) in enumerate(t3_data):
    row = table3.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[0].width = Cm(8)
    row.cells[1].width = Cm(9.5)

doc.add_paragraph()

# Подписи
sig_table = doc.add_table(rows=1, cols=2)
sig_table.style = 'Table Grid'
left = sig_table.rows[0].cells[0]
right = sig_table.rows[0].cells[1]

left.text = "Заказчик:\nТоварищество с ограниченной ответственностью\n«Astana Railway Group Trading »"
right.text = "Перевозчик:\nИП DANA CARGO SHIPPING"

for cell in [left, right]:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph()

# М.П. строка
mp_para = doc.add_paragraph()
normal_run(mp_para, "______________________ М.П.                                    _______________________ М.П.")

doc.save("template.docx")
print("template.docx создан успешно!")
