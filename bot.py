import logging
import os
from dotenv import load_dotenv
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    filters, ContextTypes, ConversationHandler
)
from docx import Document
import gspread
from google.oauth2.service_account import Credentials

load_dotenv(os.path.join(os.path.dirname(__file__), ".env.example"))

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

BOT_TOKEN = os.getenv("BOT_TOKEN")
if not BOT_TOKEN:
    raise RuntimeError("BOT_TOKEN not found in .env.example")

SPREADSHEET_ID = os.getenv("SPREADSHEET_ID")
if not SPREADSHEET_ID:
    raise RuntimeError("SPREADSHEET_ID not found in .env.example")

# --- Google Sheets ---
SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets.readonly",
    "https://www.googleapis.com/auth/drive.readonly",
]
CREDENTIALS_PATH = os.path.join(os.path.dirname(__file__), "credentials.json")
creds = Credentials.from_service_account_file(CREDENTIALS_PATH, scopes=SCOPES)
gc = gspread.authorize(creds)
sheet = gc.open_by_key(SPREADSHEET_ID).sheet1

SHEET_HEADERS = {
    "gruzopoluchatel": 1,  # B
    "marshrut": 2,         # C
    "voditel": 4,          # E
    "telefon": 5,          # F
    "iin": 6,              # G
    "gosno": 7,            # H
    "vydano": 8,           # I
    "tovary": 10,          # K
}


def load_sheet_data():
    all_values = sheet.get_all_values()
    if len(all_values) < 3:
        return []
    rows = []
    for row in all_values[2:]:
        if not any(cell.strip() for cell in row):
            continue
        def cell(idx):
            return row[idx].strip() if idx < len(row) else ""
        rows.append({
            "Грузополучатель": cell(SHEET_HEADERS["gruzopoluchatel"]),
            "Маршрут перевозки": cell(SHEET_HEADERS["marshrut"]),
            "Водитель": cell(SHEET_HEADERS["voditel"]),
            "Телефон": cell(SHEET_HEADERS["telefon"]),
            "ИИН": cell(SHEET_HEADERS["iin"]),
            "Госномер машины": cell(SHEET_HEADERS["gosno"]),
            "Выдано": cell(SHEET_HEADERS["vydano"]),
            "Товары": cell(SHEET_HEADERS["tovary"]),
        })
    return rows


def get_gruzopoluchateli():
    data = load_sheet_data()
    names, seen = [], set()
    for row in data:
        name = row["Грузополучатель"]
        if name and name not in seen:
            names.append(name)
            seen.add(name)
    return names


def get_marshruty(gruzopoluchatel: str):
    data = load_sheet_data()
    routes, seen = [], set()
    for row in data:
        if row["Грузополучатель"] == gruzopoluchatel:
            route = row["Маршрут перевозки"]
            if route and route not in seen:
                routes.append(route)
                seen.add(route)
    return routes


def get_tovary():
    data = load_sheet_data()
    items, seen = [], set()
    for row in data:
        t = row["Товары"]
        if t and t not in seen:
            items.append(t)
            seen.add(t)
    return items


def get_voditeli():
    """Get unique drivers with their info."""
    data = load_sheet_data()
    drivers = {}
    for row in data:
        name = row["Водитель"]
        if name and name not in drivers:
            drivers[name] = {
                "fio": name,
                "iin": row["ИИН"],
                "gosno": row["Госномер машины"],
                "vydano": row["Выдано"],
                "tel": row["Телефон"],
            }
    return drivers


# --- Состояния диалога ---
(
    ZAYAVKA_NUM, DATE,
    GRUZOPOLUCHATEL, MARSHRUT, TIP_GRUZA, TEMPERATURA,
    STOIMOST,
    VODITEL_SELECT, VODITEL_CONFIRM, VODITEL_EDIT_FIELD, VODITEL_EDIT_VALUE,
    VODITEL_MANUAL_FIO, VODITEL_MANUAL_IIN, VODITEL_MANUAL_GOSNO,
    VODITEL_MANUAL_VYDANO, VODITEL_MANUAL_TEL,
    CONFIRM
) = range(17)

SKIP = "\u2014"

# Поля заявки (для итогового резюме и документа)
F_ZAYAVKA_NUM = "zayavka_num"
F_DATE = "date"
F_GRUZOPOLUCHATEL = "gruzopoluchatel"
F_MARSHRUT = "marshrut"
F_TIP_GRUZA = "tip_gruza"
F_TEMPERATURA = "temperatura"
F_STOIMOST = "stoimost"
F_FIO = "fio_voditel"
F_IIN = "iin"
F_GOSNO = "gosno"
F_VYDANO = "vydano"
F_TEL = "tel_voditel"

FIELD_LABELS = {
    F_ZAYAVKA_NUM: "Номер заявки",
    F_DATE: "Дата",
    F_GRUZOPOLUCHATEL: "Грузополучатель",
    F_MARSHRUT: "Маршрут перевозки",
    F_TIP_GRUZA: "Тип груза",
    F_TEMPERATURA: "Температура",
    F_STOIMOST: "Стоимость перевозчика",
    F_FIO: "Ф.И.О. водителя",
    F_IIN: "ИИН водителя",
    F_GOSNO: "Госномер машины",
    F_VYDANO: "Выдано МВД РК от",
    F_TEL: "Тел. водителя",
}

SUMMARY_ORDER = [
    F_ZAYAVKA_NUM, F_DATE, F_GRUZOPOLUCHATEL, F_MARSHRUT,
    F_TIP_GRUZA, F_TEMPERATURA, F_STOIMOST,
    F_FIO, F_IIN, F_GOSNO, F_VYDANO, F_TEL,
]

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "new_template.docx")

# Маппинг полей водителя для кнопок редактирования
DRIVER_EDIT_MAP = {
    "Изменить госномер": F_GOSNO,
    "Изменить ИИН": F_IIN,
    "Изменить телефон": F_TEL,
    "Изменить выдано": F_VYDANO,
    "Изменить ФИО": F_FIO,
}


# --- Генерация документа ---

def _fill(cell, text):
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    para = cell.paragraphs[0]
    if para.runs:
        font_size = para.runs[0].font.size
        font_name = para.runs[0].font.name
        for run in para.runs:
            run.text = ""
        para.runs[0].text = text
        para.runs[0].font.size = font_size
        if font_name:
            para.runs[0].font.name = font_name
    else:
        para.add_run(text)


def generate_document(data: dict) -> str:
    doc = Document(TEMPLATE_PATH)
    num = data.get(F_ZAYAVKA_NUM, "")
    date = data.get(F_DATE, "")

    for para in doc.paragraphs:
        text = para.text
        if "Заявка" in text and "№" in text:
            for run in para.runs:
                if "№" in run.text or any(c.isdigit() for c in run.text):
                    run.text = run.text.replace("03", num).replace("№03", f"№{num}")
            if "03" in para.text:
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = f"Заявка №{num}"
        if "г. Астана" in text and "Марта" in text:
            # Runs: "г." " " "Астана" "\t18" " " "Марта" " " "2026" " г."
            # Replace date runs (from \t18 onwards) with single date value
            found_date_start = False
            first_date_run = None
            for run in para.runs:
                if "\t" in run.text and any(c.isdigit() for c in run.text):
                    found_date_start = True
                    first_date_run = run
                    run.text = f"\t{date}"
                elif found_date_start:
                    run.text = ""

    t0 = doc.tables[0]
    _fill(t0.rows[2].cells[1], data.get(F_GRUZOPOLUCHATEL, ""))
    _fill(t0.rows[3].cells[1], data.get(F_MARSHRUT, ""))
    _fill(t0.rows[4].cells[1], "Газель")
    _fill(t0.rows[5].cells[1], data.get(F_TIP_GRUZA, ""))
    _fill(t0.rows[6].cells[1], data.get(F_TEMPERATURA, ""))

    t1 = doc.tables[1]
    _fill(t1.rows[0].cells[1], data.get(F_STOIMOST, ""))
    gosno = data.get(F_GOSNO, "")
    _fill(t1.rows[1].cells[1], f"Газель, {gosno}" if gosno and gosno != SKIP else "Газель")
    fio = data.get(F_FIO, "")
    iin = data.get(F_IIN, "")
    vydano = data.get(F_VYDANO, "")
    tel = data.get(F_TEL, "")
    _fill(t1.rows[2].cells[1], f"{fio}\nИИН {iin}\nВыдано МВД РК от {vydano}\nТел.: {tel}")

    # Фиксируем позиции печатей — меняем с paragraph на page absolute
    _fix_stamp_positions(doc)

    safe_num = "".join(c for c in num if c.isalnum() or c in "-_")
    output_path = os.path.join(os.path.dirname(__file__), f"zayavka_{safe_num}.docx")
    doc.save(output_path)
    return output_path


WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MP_LINE_Y = 9300000  # приблизительная позиция строки "м.п." от верха A4 страницы (EMU)


def _fix_stamp_positions(doc):
    """Меняем anchor печатей с relativeFrom=paragraph на relativeFrom=page."""
    for para in doc.paragraphs:
        if "м.п." not in para.text:
            continue
        drawings = para._element.findall(f".//{{{W_NS}}}drawing")
        for drawing in drawings:
            for anchor in drawing.findall(f"{{{WP_NS}}}anchor"):
                posV = anchor.find(f"{{{WP_NS}}}positionV")
                if posV is None:
                    continue
                if posV.get("relativeFrom") == "page":
                    continue
                offset_el = posV.find(f"{{{WP_NS}}}posOffset")
                if offset_el is None:
                    continue
                old_offset = int(offset_el.text)
                posV.set("relativeFrom", "page")
                offset_el.text = str(MP_LINE_Y + old_offset)


# --- Хелперы ---

def answers(context):
    return context.user_data.setdefault("answers", {})


def build_summary(data: dict) -> str:
    lines = ["*Проверьте данные заявки:*\n"]
    for key in SUMMARY_ORDER:
        label = FIELD_LABELS[key]
        val = data.get(key, "")
        if val and val != SKIP:
            lines.append(f"*{label}:* {val}")
    lines.append("\nВсё верно?")
    return "\n".join(lines)


def driver_summary(data: dict) -> str:
    lines = ["*Данные водителя:*\n"]
    lines.append(f"*ФИО:* {data.get(F_FIO, '')}")
    lines.append(f"*ИИН:* {data.get(F_IIN, '')}")
    lines.append(f"*Госномер:* {data.get(F_GOSNO, '')}")
    lines.append(f"*Выдано:* {data.get(F_VYDANO, '')}")
    lines.append(f"*Телефон:* {data.get(F_TEL, '')}")
    return "\n".join(lines)


# --- Обработчики ---

async def start_first_question(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data.clear()
    context.user_data["answers"] = {}
    await update.message.reply_text(
        "Заполняем заявку на перевозку.\n\n"
        "Для необязательных полей можно написать *—* чтобы пропустить.\n"
        "/cancel — отмена в любой момент.\n\n"
        "*Номер заявки?*",
        parse_mode="Markdown",
        reply_markup=ReplyKeyboardRemove()
    )
    return ZAYAVKA_NUM


async def on_zayavka_num(update: Update, context: ContextTypes.DEFAULT_TYPE):
    answers(context)[F_ZAYAVKA_NUM] = update.message.text.strip()
    await update.message.reply_text(
        "*Дата заявки?* (например: 18 Марта 2026 г.)",
        parse_mode="Markdown"
    )
    return DATE


async def on_date(update: Update, context: ContextTypes.DEFAULT_TYPE):
    answers(context)[F_DATE] = update.message.text.strip()
    markup = ReplyKeyboardRemove()
    try:
        items = get_gruzopoluchateli()
        if items:
            markup = ReplyKeyboardMarkup([[n] for n in items], resize_keyboard=True, one_time_keyboard=True)
    except Exception as e:
        logger.warning(f"Sheet error: {e}")
    await update.message.reply_text("*Грузополучатель?*", parse_mode="Markdown", reply_markup=markup)
    return GRUZOPOLUCHATEL


async def on_gruzopoluchatel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    answers(context)[F_GRUZOPOLUCHATEL] = update.message.text.strip()
    markup = ReplyKeyboardRemove()
    try:
        routes = get_marshruty(answers(context)[F_GRUZOPOLUCHATEL])
        if routes:
            markup = ReplyKeyboardMarkup([[r] for r in routes], resize_keyboard=True, one_time_keyboard=True)
    except Exception as e:
        logger.warning(f"Sheet error: {e}")
    await update.message.reply_text("*Маршрут перевозки?*", parse_mode="Markdown", reply_markup=markup)
    return MARSHRUT


async def on_marshrut(update: Update, context: ContextTypes.DEFAULT_TYPE):
    answers(context)[F_MARSHRUT] = update.message.text.strip()
    markup = ReplyKeyboardRemove()
    try:
        tovary = get_tovary()
        if tovary:
            markup = ReplyKeyboardMarkup([[t] for t in tovary], resize_keyboard=True, one_time_keyboard=True)
    except Exception as e:
        logger.warning(f"Sheet error: {e}")
    await update.message.reply_text("*Тип груза?*", parse_mode="Markdown", reply_markup=markup)
    return TIP_GRUZA


async def on_tip_gruza(update: Update, context: ContextTypes.DEFAULT_TYPE):
    answers(context)[F_TIP_GRUZA] = update.message.text.strip()
    await update.message.reply_text(
        "*Температура для перевозки?* _(можно пропустить — напиши —)_",
        parse_mode="Markdown", reply_markup=ReplyKeyboardRemove()
    )
    return TEMPERATURA


async def on_temperatura(update: Update, context: ContextTypes.DEFAULT_TYPE):
    val = update.message.text.strip()
    if val != SKIP:
        answers(context)[F_TEMPERATURA] = val
    await update.message.reply_text(
        "*Стоимость услуг Перевозчика?* _(можно пропустить — напиши —)_",
        parse_mode="Markdown"
    )
    return STOIMOST


async def on_stoimost(update: Update, context: ContextTypes.DEFAULT_TYPE):
    val = update.message.text.strip()
    if val != SKIP:
        answers(context)[F_STOIMOST] = val
    # Показываем выбор водителя
    markup = ReplyKeyboardRemove()
    try:
        drivers = get_voditeli()
        if drivers:
            keyboard = [[name] for name in drivers.keys()]
            keyboard.append(["Другой водитель"])
            markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True, one_time_keyboard=True)
    except Exception as e:
        logger.warning(f"Sheet error: {e}")
    await update.message.reply_text("*Выберите водителя:*", parse_mode="Markdown", reply_markup=markup)
    return VODITEL_SELECT


async def on_voditel_select(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()

    if text == "Другой водитель":
        await update.message.reply_text(
            "*Ф.И.О. водителя?*", parse_mode="Markdown", reply_markup=ReplyKeyboardRemove()
        )
        return VODITEL_MANUAL_FIO

    # Ищем водителя в базе
    try:
        drivers = get_voditeli()
        info = drivers.get(text)
    except Exception as e:
        logger.warning(f"Sheet error: {e}")
        info = None

    if not info:
        await update.message.reply_text(
            "Водитель не найден. Введите *Ф.И.О. водителя* вручную:",
            parse_mode="Markdown", reply_markup=ReplyKeyboardRemove()
        )
        return VODITEL_MANUAL_FIO

    # Заполняем данные водителя
    a = answers(context)
    a[F_FIO] = info["fio"]
    a[F_IIN] = info["iin"]
    a[F_GOSNO] = info["gosno"]
    a[F_VYDANO] = info["vydano"]
    a[F_TEL] = info["tel"]

    # Показываем данные + кнопки правки
    keyboard = [
        ["Всё верно"],
        ["Изменить госномер", "Изменить ИИН"],
        ["Изменить телефон", "Изменить выдано"],
        ["Изменить ФИО"],
    ]
    await update.message.reply_text(
        driver_summary(a) + "\n\nВсё верно или хотите что-то изменить?",
        parse_mode="Markdown",
        reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    )
    return VODITEL_CONFIRM


async def on_voditel_confirm(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()

    if text == "Всё верно":
        # Переходим к финальному подтверждению
        keyboard = [["Подтвердить", "Отмена"]]
        await update.message.reply_text(
            build_summary(answers(context)),
            parse_mode="Markdown",
            reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
        )
        return CONFIRM

    if text in DRIVER_EDIT_MAP:
        field = DRIVER_EDIT_MAP[text]
        context.user_data["editing_field"] = field
        label = FIELD_LABELS[field]
        current = answers(context).get(field, "")
        await update.message.reply_text(
            f"Текущее значение *{label}*: {current}\n\nВведите новое значение:",
            parse_mode="Markdown", reply_markup=ReplyKeyboardRemove()
        )
        return VODITEL_EDIT_VALUE

    await update.message.reply_text("Выберите действие из кнопок.")
    return VODITEL_CONFIRM


async def on_voditel_edit_value(update: Update, context: ContextTypes.DEFAULT_TYPE):
    field = context.user_data.get("editing_field")
    answers(context)[field] = update.message.text.strip()

    # Снова показываем данные водителя
    keyboard = [
        ["Всё верно"],
        ["Изменить госномер", "Изменить ИИН"],
        ["Изменить телефон", "Изменить выдано"],
        ["Изменить ФИО"],
    ]
    await update.message.reply_text(
        driver_summary(answers(context)) + "\n\nВсё верно или хотите ещё что-то изменить?",
        parse_mode="Markdown",
        reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    )
    return VODITEL_CONFIRM


# --- Ручной ввод водителя ---

async def on_manual_fio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    answers(context)[F_FIO] = update.message.text.strip()
    await update.message.reply_text("*ИИН водителя?* _(можно пропустить — напиши —)_", parse_mode="Markdown")
    return VODITEL_MANUAL_IIN


async def on_manual_iin(update: Update, context: ContextTypes.DEFAULT_TYPE):
    val = update.message.text.strip()
    if val != SKIP:
        answers(context)[F_IIN] = val
    await update.message.reply_text("*Гос. номер машины?* _(можно пропустить — напиши —)_", parse_mode="Markdown")
    return VODITEL_MANUAL_GOSNO


async def on_manual_gosno(update: Update, context: ContextTypes.DEFAULT_TYPE):
    val = update.message.text.strip()
    if val != SKIP:
        answers(context)[F_GOSNO] = val
    await update.message.reply_text("*Выдано МВД РК от (дата)?* _(можно пропустить — напиши —)_", parse_mode="Markdown")
    return VODITEL_MANUAL_VYDANO


async def on_manual_vydano(update: Update, context: ContextTypes.DEFAULT_TYPE):
    val = update.message.text.strip()
    if val != SKIP:
        answers(context)[F_VYDANO] = val
    await update.message.reply_text("*Телефон водителя?* _(можно пропустить — напиши —)_", parse_mode="Markdown")
    return VODITEL_MANUAL_TEL


async def on_manual_tel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    val = update.message.text.strip()
    if val != SKIP:
        answers(context)[F_TEL] = val
    # Финальное подтверждение
    keyboard = [["Подтвердить", "Отмена"]]
    await update.message.reply_text(
        build_summary(answers(context)),
        parse_mode="Markdown",
        reply_markup=ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    )
    return CONFIRM


# --- Подтверждение и генерация ---

async def confirm(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text.strip()
    if "Подтвердить" in text:
        await update.message.reply_text("Генерирую документ...", reply_markup=ReplyKeyboardRemove())
        try:
            path = generate_document(answers(context))
            with open(path, "rb") as f:
                await update.message.reply_document(
                    document=f,
                    filename=os.path.basename(path),
                    caption="Заявка сформирована!"
                )
            os.remove(path)
        except Exception as e:
            logger.exception(e)
            await update.message.reply_text(f"Ошибка при генерации: {e}")
        return ConversationHandler.END
    else:
        await update.message.reply_text(
            "Отменено. Напиши /start чтобы начать заново.",
            reply_markup=ReplyKeyboardRemove()
        )
        return ConversationHandler.END


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Отменено. Напиши /start чтобы начать заново.",
        reply_markup=ReplyKeyboardRemove()
    )
    return ConversationHandler.END


def main():
    app = Application.builder().token(BOT_TOKEN).build()
    tf = filters.TEXT & ~filters.COMMAND

    conv = ConversationHandler(
        entry_points=[CommandHandler("start", start_first_question)],
        states={
            ZAYAVKA_NUM:        [MessageHandler(tf, on_zayavka_num)],
            DATE:               [MessageHandler(tf, on_date)],
            GRUZOPOLUCHATEL:    [MessageHandler(tf, on_gruzopoluchatel)],
            MARSHRUT:           [MessageHandler(tf, on_marshrut)],
            TIP_GRUZA:          [MessageHandler(tf, on_tip_gruza)],
            TEMPERATURA:        [MessageHandler(tf, on_temperatura)],
            STOIMOST:           [MessageHandler(tf, on_stoimost)],
            VODITEL_SELECT:     [MessageHandler(tf, on_voditel_select)],
            VODITEL_CONFIRM:    [MessageHandler(tf, on_voditel_confirm)],
            VODITEL_EDIT_VALUE: [MessageHandler(tf, on_voditel_edit_value)],
            VODITEL_MANUAL_FIO:   [MessageHandler(tf, on_manual_fio)],
            VODITEL_MANUAL_IIN:   [MessageHandler(tf, on_manual_iin)],
            VODITEL_MANUAL_GOSNO: [MessageHandler(tf, on_manual_gosno)],
            VODITEL_MANUAL_VYDANO:[MessageHandler(tf, on_manual_vydano)],
            VODITEL_MANUAL_TEL:   [MessageHandler(tf, on_manual_tel)],
            CONFIRM:            [MessageHandler(tf, confirm)],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )

    app.add_handler(conv)
    print("Бот запущен...")
    app.run_polling()


if __name__ == "__main__":
    main()
